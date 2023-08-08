# -*- coding: utf-8 -*-
import random,os,json,io
import pandas as pd
import streamlit as st
import  streamlit_toggle as tog
from langchain.llms import OpenAI
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.embeddings import HuggingFaceEmbeddings
from langchain.vectorstores import FAISS
import openai 
import fitz
import docx
from gtts import gTTS
import PyPDF2
from PyPDF2 import PdfReader
from utils import text_to_docs
from langchain import PromptTemplate, LLMChain
from langchain.chat_models import ChatOpenAI
from langchain.chains import ConversationChain
from langchain.memory import ConversationBufferMemory
from langchain.memory import ConversationSummaryBufferMemory
from langchain.chains.conversation.prompt import ENTITY_MEMORY_CONVERSATION_TEMPLATE
from langchain.chains.conversation.memory import ConversationEntityMemory
from langchain.callbacks import get_openai_callback
from io import StringIO
from io import BytesIO
from usellm import Message, Options, UseLLM
#from playsound import playsound
#from langchain.text_splitter import CharacterTextSplitter
#from langchain.embeddings.openai import OpenAIEmbeddings
#from langchain.chains.summarize import load_summarize_chain
#import os
#import pyaudio
#import wave
#from langchain.document_loaders import UnstructuredPDFLoader
#import streamlit.components.v1 as components
#from st_custom_components import st_audiorec, text_to_docs
#import sounddevice as sd
#from scipy.io.wavfile import write

# Setting Env
os.environ["OPENAI_API_KEY"] = st.secrets["OPENAI_API_KEY"]


@st.cache_data
def merge_pdfs(pdf_list):
    """
    Helper function to merge PDFs
    """
    pdf_merger = PyPDF2.PdfMerger()
    for pdf in pdf_list:
        pdf_document = PyPDF2.PdfReader(pdf)
        pdf_merger.append(pdf_document)
    output_pdf = BytesIO()
    pdf_merger.write(output_pdf)
    pdf_merger.close()
    return output_pdf


@st.cache_data
def merge_and_extract_text(pdf_list):
    """
    Helper function to merge PDFs and extract text
    """
    pdf_merger = PyPDF2.PdfMerger()
    for pdf in pdf_list:
        with open(pdf, 'rb') as file:
            pdf_merger.append(file)
    output_pdf = BytesIO()
    pdf_merger.write(output_pdf)
    pdf_merger.close()
    
    # Extract text from merged PDF
    merged_pdf = PyPDF2.PdfReader(output_pdf)
    all_text = []
    for page in merged_pdf.pages:
        text = page.extract_text()
        all_text.append(text)
    
    return ' '.join(all_text)

def reset_session_state():
    session_state = st.session_state
    session_state.clear()

# def merge_and_extract_text(pdf_list):
#     merged_pdf = fitz.open()
#     # Merge the PDF files
#     for pdf_file in pdf_list:
#         pdf_document = fitz.open(pdf_file)
#         merged_pdf.insert_pdf(pdf_document)
#     # Create an empty string to store the extracted text
#     merged_text = ""
#     # Extract text from each page of the merged PDF
#     for page_num in range(merged_pdf.page_count):
#         page = merged_pdf[page_num]
#         text = page.get_text()
#         merged_text += text
#     # Close the merged PDF
#     merged_pdf.close()
#     return merged_text


@st.cache_data
def render_pdf_as_images(pdf_file):
    """
    Helper function to render PDF pages as images
    """
    pdf_images = []
    pdf_document = fitz.open(stream=pdf_file.read(), filetype="pdf")
    for page_num in range(pdf_document.page_count):
        page = pdf_document.load_page(page_num)
        img = page.get_pixmap()
        img_bytes = img.tobytes()
        pdf_images.append(img_bytes)
    pdf_document.close()
    return pdf_images

# Setting globals
if "visibility" not in st.session_state:
    st.session_state.visibility = "visible"
    st.session_state.disabled = True
if "stored_session" not in st.session_state:
    st.session_state["stored_session"] = []
if "tmp_table" not in st.session_state:
    st.session_state.tmp_table=pd.DataFrame()
if "tmp_summary" not in st.session_state:
    st.session_state["tmp_summary"] = ''
if "case_num" not in st.session_state:
    st.session_state.case_num = ''

# Apply CSS styling to resize the buttons
st.markdown("""
    <style>
        .stButton button {
            width: 200px;
            height: 50px;
        }
    </style>
""", unsafe_allow_html=True)

# Applying to download button -> download_button
st.markdown("""
    <style>
        .stButton download_button {
            width: 200px;
            height: 50px;
        }
    </style>
""", unsafe_allow_html=True)

# Set Sidebar
st.markdown("""
<style>
    [data-testid=stSidebar] {
        background-color: FBFBFB;
    }
</style>
""", unsafe_allow_html=True)

st.title("Suspicious Activity Reporting Assistant")
st.markdown('---')
with st.sidebar:
    # st.sidebar.write("This is :blue[test]")
    # Navbar
    st.markdown('<link rel="stylesheet" href="https://maxcdn.bootstrapcdn.com/bootstrap/4.0.0/css/bootstrap.min.css" integrity="sha384-Gn5384xqQ1aoWXA+058RXPxPg6fy4IWvTNh0E263XmFcJlSAwiGgFAW/dAiS6JXm" crossorigin="anonymous">', unsafe_allow_html=True)

    st.markdown("""
    <nav class="navbar fixed-top navbar-expand-lg navbar-dark" style="background-color: #000000;">
    <button class="navbar-toggler" type="button" data-toggle="collapse" data-target="#navbarNav" aria-controls="navbarNav" aria-expanded="false" aria-label="Toggle navigation">
        <span class="navbar-toggler-icon"></span>
    </button>
    <div class="collapse navbar-collapse" id="navbarNav">
        <ul class="navbar-nav">
        <li class="nav-item active">
            <a class="nav-link disabled" href="#">
                <img src="	https://www.exlservice.com/themes/exl_service/exl_logo_rgb_orange_pos_94.png" width="50" height="30" alt="">
                <span class="sr-only">(current)</span>
            </a>
        </li>
        </ul>
    </div>
    </nav>
    """, unsafe_allow_html=True)

    # Add the app name
    st.sidebar.markdown('<p class="big-font">SARA</p>', unsafe_allow_html=True)
    # st.sidebar.header("SARA")
    # Add a single dropdown
    options = ["Select a Case ID", "SAR-2023-24680", "SAR-2023-13579", "SAR-2023-97531", "SAR-2023-86420", "SAR-2023-24681"]
    selected_option = st.sidebar.selectbox("", options)
    # Add the image to the sidebar below options
    st.sidebar.image("MicrosoftTeams-image (7).png", use_column_width=True)

# Redirect to Merge PDFs page when "Merge PDFs" is selected
if selected_option == "SAR-2023-24680":
    st.session_state.case_num = "SAR-2023-24680"
    # st.header("Merge Documents")
    # st.write("Upload multiple document files and merge them into one doc.")

    # Upload PDF files
    # st.subheader("Upload Case Files")
    st.write("Case No: ",st.session_state.case_num)
    pdf_files = st.file_uploader("Upload Evidences", type=["pdf"], accept_multiple_files=True)

    # Show uploaded files in a dropdown
    if pdf_files:
        st.subheader("Uploaded Files")
        file_names = [file.name for file in pdf_files]
        selected_file = st.selectbox("Select a file", file_names)
        # Enabling the button
        st.session_state.disabled = False
        # Display selected PDF contents
        if selected_file:
            selected_pdf = [pdf for pdf in pdf_files if pdf.name == selected_file][0]
            pdf_images = render_pdf_as_images(selected_pdf)
            st.subheader(f"Contents of {selected_file}")
            for img_bytes in pdf_images:
                st.image(img_bytes, use_column_width=True)
else:
    # Disabling the button
    st.session_state.disabled = True
    st.session_state.case_num = selected_option


model_name = "sentence-transformers/all-MiniLM-L6-v2"
# model_name = "hkunlp/instructor-large"

# Memory setup
llm = ChatOpenAI(temperature=0.0)
memory = ConversationSummaryBufferMemory(llm=llm, max_token_limit=500)
conversation = ConversationChain(llm=llm, memory =memory,verbose=False)


@st.cache_data
def usellm(prompt):

    service = UseLLM(service_url="https://usellm.org/api/llm")
    messages = [
      Message(role="system", content="You are a fraud analyst, who is an expert at finding out suspicious activities"),
      Message(role="user", content=f"{prompt}"),
      ]
    options = Options(messages=messages)
    response = service.chat(options)
    return response.content


@st.cache_resource
def embed(model_name):
    hf_embeddings = HuggingFaceEmbeddings(model_name=model_name)
    return hf_embeddings

hf_embeddings = embed(model_name) 


# Vizualising the files
# st.header("Welcome to the PDF Merger App")
# st.write("Use the navigation sidebar to merge PDF files.")

# # Add a single dropdown
# st.sidebar.markdown(" ")
# st.sidebar.markdown("---")
# options = ["Select a Case", "Case 1", "Case 2", "Case 3", "Case 4", "Case 5"]
# selected_option = st.sidebar.selectbox("Select a Case No.", options)

# # Redirect to Merge PDFs page when "Merge PDFs" is selected
# if selected_option == "Case 1":
#     # st.header("Merge Documents")
#     # st.write("Upload multiple document files and merge them into one doc.")

#     # Upload PDF files
#     # st.subheader("Upload Case Files")
#     pdf_files = st.file_uploader("Choose files", type=["pdf"], accept_multiple_files=True)
    
#     # Show uploaded files in a dropdown
#     if pdf_files:
#         st.subheader("Uploaded Files:")
#         file_names = [file.name for file in pdf_files]
#         selected_file = st.selectbox("Select a file", file_names)

#         # Display selected PDF contents
#         if selected_file:
#             selected_pdf = [pdf for pdf in pdf_files if pdf.name == selected_file][0]
#             pdf_images = render_pdf_as_images(selected_pdf)
#             st.subheader(f"Contents of {selected_file}")
#             for img_bytes in pdf_images:
#                 st.image(img_bytes, use_column_width=True)              
                
    # Merge PDFs extract text
    # if st.button("Merge and Download"):
    #     if pdf_files:
    #         merged_pdf = merge_pdfs(pdf_files)
    #         # Extract text from merged PDF
    #         final_pdf = PyPDF2.PdfReader(merged_pdf)
    #         all_text = []
    #         global final_txt
    #         for page in final_pdf.pages:
    #             text = page.extract_text()
    #             all_text.append(text)
    #         final_txt = ' '.join(all_text)

            # downloading content
            # st.download_button(
            #     label="Download Merged PDF",
            #     data=merged_pdf.getvalue(),
            #     file_name="merged_pdf.pdf",
            #     mime="application/pdf",
            # )



text_splitter = RecursiveCharacterTextSplitter(
    chunk_size = 1000,
    chunk_overlap  = 100,
    length_function = len,
    separators=["\n\n", "\n", " ", ""]
)
#text_splitter = CharacterTextSplitter.from_tiktoken_encoder(chunk_size=100, chunk_overlap=0)
#texts = ''

# @st.cache_data
# def embedding_store(file):
#     # save file
#     pdf_reader = PdfReader(file)
#     text = ""
#     for page in pdf_reader.pages:
#         text += page.extract_text()
#     #st.write(text)
#     texts =  text_splitter.split_text(text)
#     docs = text_to_docs(texts)
#     #st.write(texts)
#     docsearch = FAISS.from_documents(docs, hf_embeddings)
#     return docs, docsearch

# Addding markdown styles(Global)
st.markdown("""
<style>
.big-font {
    font-size:60px !important;
}
</style>
""", unsafe_allow_html=True)

@st.cache_data
def embedding_store(pdf_files):
    merged_pdf = merge_pdfs(pdf_files)
    final_pdf = PyPDF2.PdfReader(merged_pdf)
    text = ""
    for page in final_pdf.pages:
        text += page.extract_text()
    texts =  text_splitter.split_text(text)
    docs = text_to_docs(texts)
    docsearch = FAISS.from_documents(docs, hf_embeddings)
    return docs, docsearch

# Creating header
col1,col2 = st.columns(2)
with col1:
    st.subheader('Pre-Set Questionnaire')
    # Create a Pandas DataFrame with your data

    data = {'Questions': [" What is the Victim's Name?",' Has any suspect been reported?',' List the Merchant name',' How was the bank notified?',' When was the bank notified?',' What is the Fraud Type?',' When did the fraud occur?',' Was the disputed amount greater than 5000 USD?',' What type of cards are involved?',' Was the police report filed?']}
    df_fixed = pd.DataFrame(data)
    df_fixed.index = df_fixed.index +1
with col2:
    # Create a checkbox to show/hide the table
    cols1, cols2, cols3, cols4 = st.columns([1,1,1,1])
    with cols1:
        show_table = tog.st_toggle_switch(label="", 
                            key="Key1", 
                            default_value=False, 
                            label_after = False, 
                            inactive_color = '#D3D3D3', 
                            active_color="#11567f", 
                            track_color="#29B5E8"
                            )
    # Show the table if the checkbox is ticked
    if show_table:
        # st.write(df_fixed)
        st.dataframe(df_fixed, width=1000)

with st.spinner('Wait for it...'):
    if st.button("Generate Insights",disabled=st.session_state.disabled):
        if pdf_files is not None:
            # File handling logic
            _, docsearch = embedding_store(pdf_files)
            queries ="Please provide the following information regarding the possible fraud case: What is the name of the customer name,\
            has any suspect been reported, list the merchant name, how was the bank notified, when was the bank notified, what is the fraud type,\
            when did the fraud occur, was the disputed amount greater than 5000 USD, what type of cards are involved, was the police report filed,\
            and based on the evidence, is this a suspicious activity(Summarize all the questions asked prior to this in a detailed manner),that's the answer of\
            whether this is a suspicious activity\
            "
            contexts = docsearch.similarity_search(queries, k=5) 
            prompts = f" Give a the answer to the below questions as truthfully and in as detailed in the form of sentences\
            as possible as per given context only,\n\n\
                    1. What is the Victim's Name?\n\
                    2. Has any suspect been reported?\n\
                    3. List the Merchant name\n\
                    4. How was the bank notified?\n\
                    5. When was the bank notified?\n\
                    6. What is the Fraud Type?\n\
                    7. When did the fraud occur?\n\
                    8. Was the disputed amount greater than 5000 USD?\n\
                    9. What type of cards are involved?\n\
                    10. Was the police report filed?\n\
                Context: {contexts}\n\
                Response (in the python dictionary format\
                where the dictionary key would carry the questions and its value would have a descriptive answer to the questions asked): "
                

            response = usellm(prompts)
            # st.write(response)
            # memory.save_context({"input": f"{queries}"}, {"output": f"{response}"})
            # st.write(response)
            # st.write(memory.load_memory_variables({}))



            # Convert the response in dictionary from tbl
            # prompt_conv = f" Convert the tabular data into a python dictionary\
            #     context: {response}\
            #     Response (give me the response in the form of a python dictionary with questions exactly as it is): "
            # resp_dict = usellm(prompt_conv)
            # st.write(resp_dict)
            resp_dict_obj = json.loads(response)
            res_df = pd.DataFrame(resp_dict_obj.items(), columns=['Question','Answer'])
            try:
                res_df.Question = res_df.Question.apply(lambda x: x.split(".")[1])
                res_df.index = res_df.index + 1
            except IndexError:
                pass
            st.table(res_df)
            # st.write(resp_dict_obj)
            st.session_state["tmp_table"] = pd.concat([st.session_state.tmp_table, res_df], ignore_index=True)
st.markdown("---")

# For input box outside of template
try:
    docs, docsearch = embedding_store(pdf_files)
except Exception:
    pass


# Text Input

st.subheader("Ask Additional Questions")
query = st.text_input('Please ask below the additional case questions.',disabled=st.session_state.disabled)
text_dict = {}

def LLM_Response():
    llm_chain = LLMChain(prompt=prompt, llm=llm)
    response = llm_chain.run({"query":query, "context":context})
    return response

with st.spinner('Getting you information...'):      
    if query:
        # Text input handling logic
        #st.write("Text Input:")
        #st.write(text_input)

        context_1 = docsearch.similarity_search(query, k=5)

        if query.lower() == "what is the victim's name?":
            prompt_1 = f'''Perform Name Enitity Recognition to identify the Customer name as accurately as possible, given the context. The Customer can also be referenced as the Victim or the person with whom the Fraud has taken place.\n\n\
                        Question: {query}\n\
                        Context: {context_1}\n\
                        Response: '''

            
        elif query.lower() == "what is the suspect's name?":
            prompt_1 = f'''Perform Name Enitity Recognition to identify the Suspect name as accurately as possible, given the context. Suspect is the Person who has committed the fraud with the Customer. Respond saying "The Suspect Name is not Present" if there is no suspect in the given context.\n\n\
                        Question: {query}\n\
                        Context: {context_1}\n\
                        Response: '''

            
        elif query.lower() == "list the merchant name":
            prompt_1 = f'''Perform Name Enitity Recognition to identify all the Merchant Organizations as accurately as possible, given the context. A merchant is a type of business or organization that accepts payments from the customer account. Give a relevant and concise response.\n\n\
                        Question: {query}\n\
                        Context: {context_1}\n\
                        Response: '''

            
        elif query.lower() == "how was the bank notified?":
            prompt_1 = f''' You need to act as a Financial analyst to identify how was the bank notified of the Supicious or Fraud event with in the given context. The means of communication can be a call, an email or in person. Give a relevant and concise response.\n\n\
                        Question: {query}\n\
                        Context: {context_1}\n\
                        Response: '''

            
        elif query.lower() == "when was the bank notified?":
            prompt_1 = f''' You need to act as a Financial analyst to identify the when the bank was notified of the Fraud i.e., the disputed date. Given the context, provide a relevant and concise response.\n\n\
                        Question: {query}\n\
                        Context: {context_1}\n\
                        Response: '''

            
        elif query.lower() == "what type of fraud is taking place?":
            prompt_1 = f''' You need to act as a Financial analyst to identify the type of fraud or suspicious activity has taken place amd summarize it, within the given context. Also mention the exact fraud code. Give a relevant and concise response.\n\n\
                        Question: {query}\n\
                        Context: {context_1}\n\
                        Response: '''

            
        elif query.lower() == "when did the fraud occur?":
            prompt_1 = f''' You need to act as a Financial analyst to identify the when the did the fraud occur i.e., the Transaction Date. Given the context, provide a relevant and concise response.\n\n\
                        Question: {query}\n\
                        Context: {context_1}\n\
                        Response: '''

                
        elif query.lower() == "was the disputed amount greater than 5000 usd?":
            prompt_1 = f''' You need to act as a Financial analyst to identify the disputed amount and perform a mathematical calculation to check if the disputed amount is greater than 5000 or no, given the context. Give a relevant and concise response.\n\n\
                        Question: {query}\n\
                        Context: {context_1}\n\
                        Response: '''

            
        elif query.lower() == "what type of cards are involved?":
            prompt_1 = f''' You need to act as a Financial analyst to identify the type of card and card's brand involved, given the context. On a higher level the card can be a Credit or Debit Card. VISA, MasterCard or American Express, Citi Group, etc. are the different brands with respect to a Credit card or Debit Card . Give a relevant and concise response.\n\n\
                        Question: {query}\n\
                        Context: {context_1}\n\
                        Response: '''

            
        elif query.lower() == "was the police report filed?":
            prompt_1 = f''' You need to act as a Financial analyst to identify if the police was reported of the Fraud activity, given the context. Give a relevant and concise response.\n\n\
                        Question: {query}\n\
                        Context: {context_1}\n\
                        Response: '''

                
        elif query.lower() == "Is this a valid SAR case?":
            prompt_1 = f''' You need to act as a Financial analyst to check if this is a SAR or not, given the following context, if the transaction amount is less than 5000 USD we cannot categorize this as SAR (Suspicious activity Report).Give a relevant and concise response. \n\n\
                        Question: {query}\n\
                        Context: {context_1}\n\
                        Response: '''

            
        else:
            prompt_1 = f'''Act as a financial analyst and give concise answer to below Question as truthfully as possible, with given Context.\n\n\
                        Question: {query}\n\
                        Context: {context_1}\n\                      
                        Response: '''


        #prompt = PromptTemplate(template=prompt, input_variables=["query", "context"])
        response = usellm(prompt_1) #LLM_Response()
        text_dict[query] = response
        # resp_dict_obj.update(text_dict)
        st.write(response)
        if response:
            df = pd.DataFrame(text_dict.items(), columns=['Question','Answer'])
        else:
            df = pd.DataFrame()
        # st.session_state.tmp_table = pd.concat([tmp_table, tmp_table], ignore_index=True)
        # st.write(text_dict.items())
        st.session_state["tmp_table"] = pd.concat([st.session_state.tmp_table, df], ignore_index=True)
        st.session_state.tmp_table.drop_duplicates(subset=['Question'])
        # st.table(st.session_state.tmp_table)


with st.spinner('Summarization ...'):  
    if st.button("Summarize",disabled=st.session_state.disabled):
        summ_dict = st.session_state.tmp_table.set_index('Question')['Answer'].to_dict()
        # chat_history = resp_dict_obj['Summary']
        memory = ConversationSummaryBufferMemory(llm=llm, max_token_limit=300)
        memory.save_context({"input": "This is the entire chat summary"}, 
                        {"output": f"{summ_dict}"})
        conversation = ConversationChain(
        llm=llm, 
        memory = memory,
        verbose=True)
        st.session_state["tmp_summary"] = conversation.predict(input="Give me a detailed summary of the above texts.")
        st.write(st.session_state["tmp_summary"] )

with st.spinner("Downloading...."):
# if st.button("Download Response", disabled=st.session_state.disabled):
    # Create a Word document with the table and some text
    doc = docx.Document()
    doc.add_paragraph(st.session_state["tmp_summary"])

    columns = list(st.session_state.tmp_table.columns)
    table = doc.add_table(rows=1, cols=len(columns), style="Table Grid")
    table.autofit = True
    for col in range(len(columns)):
        # set_cell_margins(table.cell(0, col), top=100, start=100, bottom=100, end=50) # set cell margin
        table.cell(0, col).text = columns[col]
    # doc.add_table(st.session_state.tmp_table.shape[0]+1, st.session_state.tmp_table.shape[1], style='Table Grid')
    
    for i, row in enumerate(st.session_state.tmp_table.itertuples()):
        table_row = table.add_row().cells # add new row to table
        for col in range(len(columns)): # iterate over each column in row and add text
            table_row[col].text = str(row[col+1]) # avoid index by adding col+1
    # save document
    # output_bytes = docx.Document.save(doc, 'output.docx')
    # st.download_button(label='Download Report', data=output_bytes, file_name='evidence.docx', mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

    bio = io.BytesIO()
    doc.save(bio)
    if doc:
        st.download_button(
            label="Download Report",
            data=bio.getvalue(),
            file_name="Report.docx",
            mime="docx",
            disabled=st.session_state.disabled
        )

# Adding Radio button
st.header("Make Decision")
selected_rad = st.radio("Select an option", ["Approved", "Decline", "Refer for review"])
if selected_rad in ("Approved", "Decline"):
    st.write("Thanks, Your response has been recorded!")
else:
    email_regex = r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$'
    email_id = st.text_input("Enter your email ID")
    if email_id and not re.match(email_regex, email_id):
        st.error("Please enter a valid email ID")
    else:
        st.write("Thanks the details has been sent for further review")

# Allow the user to clear all stored conversation sessions
# if st.button("Reset Session"):
#     reset_session_state()
#     pdf_files.clear()

# Footer
st.markdown(
    """
    <style>
      #MainMenu {visibility: hidden;}
      footer {visibility: hidden;}
    </style>
    """
    , unsafe_allow_html=True)
st.markdown('<div class="footer"><p></p></div>', unsafe_allow_html=True)

hide_st_style = """
            <style>
            #MainMenu {visibility: hidden;}
            footer {visibility: hidden;}
            header {visibility: hidden;}
            </style>
            """
st.markdown(hide_st_style, unsafe_allow_html=True)


padding = 0
st.markdown(f""" <style>
    .reportview-container .main .block-container{{
        padding-top: {padding}rem;
        padding-right: {padding}rem;
        padding-left: {padding}rem;
        padding-bottom: {padding}rem;
    }} </style> """, unsafe_allow_html=True)



